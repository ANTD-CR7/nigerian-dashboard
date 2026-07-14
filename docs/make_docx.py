"""Convert docs/FYP_REPORT.md to docs/FYP_REPORT.docx with python-docx.

Pragmatic converter for this report's Markdown subset: headings, bold runs,
pipe tables, fenced code blocks (mermaid blocks are kept as monospace source
with a note to render them at mermaid.live), bullet/numbered lists and
blockquotes. Re-run after editing the Markdown:  python docs/make_docx.py
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor

from PIL import Image

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = Path(__file__).parent / "FYP_REPORT.md"
DST = Path(__file__).parent / "FYP_REPORT.docx"


def add_runs(par, text):
    """Render **bold** and *italic* segments as runs; strip backticks."""
    text = text.replace("`", "")
    for i, bchunk in enumerate(re.split(r"\*\*", text)):
        if not bchunk:
            continue
        bold = (i % 2 == 1)
        for j, ichunk in enumerate(re.split(r"\*", bchunk)):
            if not ichunk:
                continue
            run = par.add_run(ichunk)
            run.bold = bold
            run.italic = (j % 2 == 1)


def add_code(doc, lines, mermaid=False):
    if mermaid:
        note = doc.add_paragraph()
        r = note.add_run("[Diagram source below — render it at mermaid.live or draw.io and replace with the image]")
        r.italic = True
        r.font.size = Pt(9)
    for line in lines:
        p = doc.add_paragraph()
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)


def add_image(doc, alt, rel_path):
    """![caption](figures/x.png) → centered picture sized to the page, + caption."""
    img = Path(__file__).parent / rel_path
    if not img.exists():
        p = doc.add_paragraph()
        r = p.add_run(f"[missing image: {rel_path}]")
        r.italic = True
        return
    w_px, h_px = Image.open(img).size
    # Chapter-4 screenshots are a figures gallery: cap their height so two
    # pack onto one page instead of stranding half-empty pages. Chapter-3
    # diagrams (ERD, use-case, sequence) stay full-size — their detail matters.
    max_h = 3.7 if "fig4_" in str(rel_path) else 7.5
    width = 6.2
    height_at_width = h_px / w_px * width
    if height_at_width > max_h:
        width = width * max_h / height_at_width
    # small diagrams shouldn't be blown up past ~2x native (96dpi assumption)
    native_in = w_px / 96
    width = min(width, max(native_in * 1.6, 3.0))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True  # never split the image from its own caption
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(str(img), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(alt)
    r.italic = True
    r.font.size = Pt(9)
    cap.paragraph_format.space_after = Pt(8)


def add_toc_field(doc):
    """Insert a live Word TOC field (headings 1-3, hyperlinked)."""
    par = doc.add_paragraph()
    run = par.add_run()
    beg = OxmlElement("w:fldChar"); beg.set(qn("w:fldCharType"), "begin")
    ins = OxmlElement("w:instrText"); ins.set(qn("xml:space"), "preserve")
    ins.text = r' TOC \o "1-3" \h \z \u '
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t"); txt.text = "Right-click here and choose Update Field to refresh the table of contents."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (beg, ins, sep, txt, end):
        run._r.append(el)


def add_table(doc, rows):
    cells = [[c.strip() for c in row.strip().strip("|").split("|")] for row in rows]
    cells = [r for r in cells if not all(re.fullmatch(r":?-{2,}:?", c or "---") for c in r)]
    if not cells:
        return
    ncols = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=ncols)
    t.style = "Table Grid"
    for ri, row in enumerate(cells):
        for ci in range(ncols):
            text = row[ci] if ci < len(row) else ""
            par = t.cell(ri, ci).paragraphs[0]
            add_runs(par, text)
            for run in par.runs:
                run.font.size = Pt(9)
                if ri == 0:
                    run.bold = True


def main():
    md = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    # ── Department template: Times New Roman 12, 1" margins ──
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)
    for name in ("Heading 1", "Heading 2", "Heading 3", "Title"):
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.font.bold = True
    doc.styles["Title"].font.size = Pt(14)
    doc.styles["Heading 1"].font.size = Pt(14)
    doc.styles["Heading 2"].font.size = Pt(12)
    doc.styles["Heading 3"].font.size = Pt(12)
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)

    current_h1 = ""
    current_h2 = ""

    def style_heading(h):
        for r in h.runs:
            r.font.name = "Times New Roman"
            r.font.color.rgb = RGBColor(0, 0, 0)
            r.bold = True
        # never orphan a heading alone at the bottom of a page
        h.paragraph_format.keep_with_next = True
        return h

    def body_format(par, center=False):
        """Template body: double-spaced, justified (centered for title page)."""
        pf = par.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY

    i = 0
    while i < len(md):
        line = md[i]

        if line.startswith("```"):
            mermaid = "mermaid" in line
            block = []
            i += 1
            while i < len(md) and not md[i].startswith("```"):
                block.append(md[i]); i += 1
            add_code(doc, block, mermaid)
            i += 1
            continue

        if line.strip().startswith("|"):
            rows = []
            while i < len(md) and md[i].strip().startswith("|"):
                rows.append(md[i]); i += 1
            add_table(doc, rows)
            doc.add_paragraph()
            continue

        m_img = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", line.strip())
        if m_img:
            add_image(doc, m_img.group(1), m_img.group(2))
            i += 1
            continue

        if line.startswith("# "):
            pass  # the Title Page section carries the title; no separate doc heading
        elif line.startswith("## "):
            current_h1 = re.sub(r"\*\*", "", line[3:]).strip()
            if current_h1 == "FRONT MATTER":
                current_h2 = ""
                i += 1
                continue
            if len(doc.element.body) > 1:  # never open the document with a blank page
                doc.add_page_break()
            current_h2 = ""
            h = style_heading(doc.add_heading(current_h1, level=1))
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("### "):
            current_h2 = re.sub(r"\*\*", "", line[4:]).strip()
            if current_h2 in ("Declaration", "Certification", "Dedication", "Acknowledgements", "Abstract", "Table of Contents", "List of Figures", "List of Tables"):
                doc.add_page_break()
            if current_h2 != "Title Page":  # sample title page has no "Title Page" label
                h = style_heading(doc.add_heading(current_h2, level=2))
                if current_h2 in ("Declaration", "Certification"):
                    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if current_h2 == "Table of Contents":
                add_toc_field(doc)
        elif line.startswith("#### "):
            style_heading(doc.add_heading(re.sub(r"\*\*", "", line[5:]).strip(), level=3))
        elif line.startswith("> "):
            if current_h1 == "":
                i += 1
                continue  # internal draft note before the first chapter — not for submission
            p = doc.add_paragraph()
            add_runs(p, line[2:])
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif re.match(r"^\s*[-*] ", line):
            item = [re.sub(r"^\s*[-*] ", "", line).strip()]
            while (i + 1 < len(md) and md[i + 1].strip()
                   and re.match(r"^\s+\S", md[i + 1])
                   and not re.match(r"^\s*[-*] ", md[i + 1])):
                i += 1
                item.append(md[i].strip())
            text = " ".join(item)
            if current_h1.startswith("REFERENCES"):
                # APA 7: no bullets, hanging indent, double-spaced
                p = doc.add_paragraph()
                add_runs(p, text)
                pf = p.paragraph_format
                pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                pf.left_indent = Inches(0.5)
                pf.first_line_indent = Inches(-0.5)
            elif current_h2 in ("List of Figures", "List of Tables"):
                # thesis convention: plain lines, no bullet markers
                p = doc.add_paragraph()
                add_runs(p, text)
                pf = p.paragraph_format
                pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p = doc.add_paragraph(style="List Bullet")
                add_runs(p, text)
                body_format(p)
        elif re.match(r"^\s*\d+\. ", line):
            item = [re.sub(r"^\s*\d+\. ", "", line).strip()]
            while (i + 1 < len(md) and md[i + 1].strip()
                   and re.match(r"^\s+\S", md[i + 1])
                   and not re.match(r"^\s*\d+\. ", md[i + 1])):
                i += 1
                item.append(md[i].strip())
            p = doc.add_paragraph(style="List Number")
            add_runs(p, " ".join(item))
            body_format(p)
        elif line.strip() in ("---", "***"):
            pass
        elif line.strip():
            block = [line.strip()]
            while (i + 1 < len(md) and md[i + 1].strip()
                   and not re.match(r"^(#{1,4} |> |\s*[-*] |\s*\d+\. |```|\||!\[)", md[i + 1])
                   and md[i + 1].strip() not in ("---", "***")):
                i += 1
                block.append(md[i].strip())
            p = doc.add_paragraph()
            add_runs(p, " ".join(block))
            on_title = current_h2 in ("Title Page", "Declaration", "Certification")
            body_format(p, center=(current_h2 == "Title Page"))
            if on_title:
                for r in p.runs:
                    r.bold = True
            # a bold "Figure X.Y — ..." caption line right before an image:
            # keep it glued to that image so the pair never splits across pages
            look = i + 1
            while look < len(md) and not md[look].strip():
                look += 1
            if look < len(md) and re.match(r"^!\[", md[look].strip()) and block[0].startswith("**"):
                p.paragraph_format.keep_with_next = True
        i += 1

    doc.save(DST)
    print(f"Wrote {DST} ({DST.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
